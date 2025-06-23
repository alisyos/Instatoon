"""
Flask 웹 애플리케이션
인스타툰 스토리보드 생성기의 웹 인터페이스를 제공합니다.
"""

from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import json
import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from main import InstaToonGenerator
from config import Config

# 글로벌 생성기 인스턴스
generator = InstaToonGenerator()

# Vercel 배포를 위한 애플리케이션 팩토리
def create_app():
    """Flask 애플리케이션 팩토리"""
    app = Flask(__name__)
    CORS(app)
    
    return app

# 애플리케이션 인스턴스 생성
app = create_app()


def storyboard_to_text(storyboard):
    """스토리보드를 읽기 쉬운 텍스트 형태로 변환합니다."""
    text_lines = []
    
    # 제목과 주제
    text_lines.append(f"📖 {storyboard['wholeTitle']}")
    text_lines.append("=" * 50)
    text_lines.append(f"📝 핵심 주제: {storyboard['storyTopic']}")
    text_lines.append("")
    
    # 해시태그
    hashtags_text = " ".join(storyboard['hashtags'])
    text_lines.append(f"🏷️ 해시태그: {hashtags_text}")
    text_lines.append("")
    text_lines.append("=" * 50)
    text_lines.append("")
    
    # 페이지별 내용
    for page in storyboard['pages']:
        text_lines.append(f"📄 페이지 {page['page']}")
        text_lines.append("-" * 30)
        
        # 등장인물
        characters = ", ".join(page['character'])
        text_lines.append(f"👥 등장인물: {characters}")
        
        # 배경
        text_lines.append(f"🎬 배경: {page['background']}")
        
        # 대사
        text_lines.append("💬 대사:")
        for char, dialogue in page['dialogue'].items():
            text_lines.append(f"   {char}: \"{dialogue}\"")
        
        # 표정/포즈
        text_lines.append(f"🎭 표정/포즈: {page['expressionPose']}")
        text_lines.append("")
    
    return "\n".join(text_lines)


def create_docx_from_storyboard(storyboard):
    """스토리보드를 DOCX 문서로 변환합니다."""
    doc = Document()
    
    # 문서 제목
    title = doc.add_heading(storyboard['wholeTitle'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 핵심 주제
    doc.add_heading('📝 핵심 주제', level=1)
    doc.add_paragraph(storyboard['storyTopic'])
    
    # 해시태그
    doc.add_heading('🏷️ 해시태그', level=1)
    hashtags_text = " ".join(storyboard['hashtags'])
    doc.add_paragraph(hashtags_text)
    
    # 페이지별 내용
    doc.add_heading('📖 스토리보드', level=1)
    
    for page in storyboard['pages']:
        # 페이지 제목
        page_heading = doc.add_heading(f"페이지 {page['page']}", level=2)
        
        # 등장인물
        doc.add_paragraph().add_run("👥 등장인물: ").bold = True
        characters = ", ".join(page['character'])
        doc.add_paragraph(characters)
        
        # 배경
        doc.add_paragraph().add_run("🎬 배경: ").bold = True
        doc.add_paragraph(page['background'])
        
        # 대사
        doc.add_paragraph().add_run("💬 대사: ").bold = True
        dialogue_para = doc.add_paragraph()
        for char, dialogue in page['dialogue'].items():
            dialogue_para.add_run(f"{char}: \"{dialogue}\"\n")
        
        # 표정/포즈
        doc.add_paragraph().add_run("🎭 표정/포즈: ").bold = True
        doc.add_paragraph(page['expressionPose'])
        
        # 페이지 구분선 (마지막 페이지가 아닌 경우)
        if page != storyboard['pages'][-1]:
            doc.add_paragraph("-" * 50)
    
    return doc


@app.route('/')
def index():
    """메인 페이지"""
    return render_template('index.html')


@app.route('/api/generate', methods=['POST'])
def generate_storyboard():
    """스토리보드 생성 API"""
    try:
        data = request.json
        
        # 입력값 검증
        if not data.get('plot'):
            return jsonify({'error': '줄거리는 필수 입력 항목입니다.'}), 400
        
        if not data.get('pages'):
            return jsonify({'error': '분량은 필수 입력 항목입니다.'}), 400
        
        # 입력값 검증
        user_input = {
            'characters': data.get('characters', ''),
            'keywords': data.get('keywords', ''),
            'plot': data.get('plot', ''),
            'pages': data.get('pages', '')
        }
        
        print(f"사용자 입력 받음: {user_input}")
        
        if not generator.validate_input(user_input['plot'], user_input['pages']):
            return jsonify({'error': '입력값이 올바르지 않습니다.'}), 400
        
        # GPT 클라이언트 상태 확인
        if not generator.gpt_client:
            return jsonify({'error': 'GPT 클라이언트가 초기화되지 않았습니다. API 키를 확인해주세요.'}), 500
        
        print("스토리보드 생성 시작...")
        
        # 스토리보드 생성
        storyboard = generator.generate_storyboard(user_input)
        
        if not storyboard:
            return jsonify({'error': '스토리보드 생성에 실패했습니다. GPT 응답을 확인해주세요.'}), 500
        
        print("스토리보드 생성 완료")
        
        # 결과 저장 (선택사항)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"storyboard_{timestamp}.json"
        
        try:
            generator.save_result(storyboard, filename)
            print(f"결과 파일 저장 완료: {filename}")
        except Exception as e:
            print(f"파일 저장 오류: {e}")
        
        # 텍스트 형태로 변환
        text_content = storyboard_to_text(storyboard)
        
        return jsonify({
            'success': True,
            'storyboard': storyboard,
            'text_content': text_content,
            'filename': filename
        })
        
    except json.JSONDecodeError as e:
        error_msg = f'잘못된 JSON 형식입니다: {str(e)}'
        print(f"JSON 디코딩 오류: {error_msg}")
        return jsonify({'error': error_msg}), 400
        
    except Exception as e:
        error_msg = f'서버 오류: {str(e)}'
        print(f"API 오류: {error_msg}")
        print(f"오류 타입: {type(e).__name__}")
        
        # 스택 트레이스 출력 (디버깅용)
        import traceback
        print(f"스택 트레이스: {traceback.format_exc()}")
        
        return jsonify({'error': error_msg}), 500


@app.route('/api/download-docx', methods=['POST'])
def download_docx():
    """DOCX 파일 다운로드"""
    try:
        data = request.json
        storyboard = data.get('storyboard')
        
        if not storyboard:
            return jsonify({'error': '스토리보드 데이터가 없습니다.'}), 400
        
        # DOCX 문서 생성
        doc = create_docx_from_storyboard(storyboard)
        
        # 메모리에 저장
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        # 파일명 생성
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"storyboard_{timestamp}.docx"
        
        return send_file(
            docx_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"DOCX 다운로드 오류: {e}")
        return jsonify({'error': f'DOCX 파일 생성 중 오류가 발생했습니다: {str(e)}'}), 500


@app.route('/api/download/<filename>')
def download_file(filename):
    """생성된 파일 다운로드"""
    try:
        if os.path.exists(filename):
            return send_file(filename, as_attachment=True)
        else:
            return jsonify({'error': '파일을 찾을 수 없습니다.'}), 404
    except Exception as e:
        return jsonify({'error': f'다운로드 오류: {str(e)}'}), 500


@app.route('/api/health')
def health_check():
    """서버 상태 확인"""
    try:
        # GPT 클라이언트 상태 확인
        gpt_status = generator.gpt_client is not None
        
        return jsonify({
            'status': 'healthy',
            'gpt_client': gpt_status,
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500


@app.errorhandler(404)
def not_found(error):
    """404 에러 핸들러"""
    return jsonify({'error': '페이지를 찾을 수 없습니다.'}), 404


@app.errorhandler(500)
def internal_error(error):
    """500 에러 핸들러"""
    return jsonify({'error': '서버 내부 오류가 발생했습니다.'}), 500